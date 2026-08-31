#!/usr/bin/env python3
"""Build and style the submission-ready DOCX from article.md."""

from __future__ import annotations

import os
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "article.md"
OUTPUT = ROOT / "artifacts" / "Artigo_SISAN_Matching.docx"

NAVY = "173A5E"
TEAL = "2C756D"
SKY = "DDEAF2"
PALE = "F4F7F9"
TEXT = "24313A"
MUTED = "5B6770"
WHITE = "FFFFFF"


def get_style(doc: Document, name: str):
    """Work around reference files whose style-name index is incomplete."""
    for style in doc.styles:
        if style.name == name:
            return style
    raise KeyError(f"Style not found: {name}")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def style_run(run, name="Aptos", size=10.5, color=TEXT, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def format_styles(doc: Document) -> None:
    normal = get_style(doc, "Normal")
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.16
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    for idx, size in ((1, 16), (2, 13), (3, 11.5)):
        style = get_style(doc, f"Heading {idx}")
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if idx < 3 else TEAL)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.space_before = Pt(14 if idx == 1 else 10)
        style.paragraph_format.space_after = Pt(5)

    existing_style_names = {style.name for style in doc.styles}
    for name in ("Title", "Subtitle", "Author"):
        if name not in existing_style_names:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    title = get_style(doc, "Title")
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True

    subtitle = get_style(doc, "Subtitle")
    subtitle.font.name = "Aptos"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor.from_string(TEAL)
    subtitle.paragraph_format.space_after = Pt(14)

    if "Caption" in existing_style_names:
        caption = get_style(doc, "Caption")
        caption.font.name = "Aptos"
        caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        caption.font.size = Pt(8.5)
        caption.font.italic = False
        caption.font.color.rgb = RGBColor.from_string(MUTED)
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(8)
        caption.paragraph_format.keep_with_next = True

    if "Block Text" in existing_style_names:
        block = get_style(doc, "Block Text")
        block.font.name = "Aptos"
        block.font.size = Pt(9.5)
        block.font.color.rgb = RGBColor.from_string(MUTED)
        block.paragraph_format.left_indent = Cm(0.6)
        block.paragraph_format.right_indent = Cm(0.6)
        block.paragraph_format.space_before = Pt(8)
        block.paragraph_format.space_after = Pt(10)

    if "Bibliography" in existing_style_names:
        bibliography = get_style(doc, "Bibliography")
        bibliography.font.name = "Aptos"
        bibliography._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        bibliography.font.size = Pt(8.8)
        bibliography.font.color.rgb = RGBColor.from_string(TEXT)
        bibliography.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bibliography.paragraph_format.line_spacing = 1.0
        bibliography.paragraph_format.space_after = Pt(3)
        bibliography.paragraph_format.widow_control = True


def format_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.35)
        section.right_margin = Cm(2.35)
        section.header_distance = Cm(0.9)
        section.footer_distance = Cm(0.85)

        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.text = "COORDENAÇÃO SEM RECURSOS?  ·  MANUSCRITO PARA REVISÃO"
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            style_run(run, size=7.5, color=MUTED, bold=True)
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), TEAL)
        borders.append(bottom)
        p_pr.append(borders)

        footer = section.footer
        footer_p = footer.paragraphs[0]
        add_page_field(footer_p)
        for run in footer_p.runs:
            style_run(run, size=8, color=MUTED)


def format_title_block(doc: Document) -> None:
    # Pandoc emits title, subtitle, author and date before the first block quote.
    title_seen = False
    for paragraph in doc.paragraphs[:10]:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name == "Title":
            title_seen = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(16)
        elif title_seen and paragraph.text.strip() == "Osmar Junior":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(1)
            for run in paragraph.runs:
                style_run(run, size=10.5, color=TEXT, bold=True)
        elif title_seen and paragraph.text.startswith("Versão para revisão manual"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(12)
            for run in paragraph.runs:
                style_run(run, size=9, color=MUTED)

    # Add a compact kicker above the title without requiring a separate cover.
    if doc.paragraphs:
        first = doc.paragraphs[0]
        kicker = OxmlElement("w:p")
        first._p.addprevious(kicker)
        p = Paragraph(kicker, first._parent)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(7)
        run = p.add_run("AVALIAÇÃO DE POLÍTICAS PÚBLICAS  /  MATCHING")
        style_run(run, size=8.5, color=TEAL, bold=True)


def format_paragraphs(doc: Document) -> None:
    in_references = False
    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = p.style.name if p.style else ""

        if style_name.startswith("Heading") and text == "Referências":
            in_references = True

        if style_name.startswith("Heading"):
            continue

        if in_references and text:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.widow_control = True
            for run in p.runs:
                style_run(run, size=8.7, color=TEXT)

        if text.startswith("Palavras-chave:") or text.startswith("Keywords:"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(10)
            for run in p.runs:
                style_run(run, size=9.5, color=MUTED)

        if text.startswith("Nota:") or text.startswith("*Nota:"):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(9)
            for run in p.runs:
                style_run(run, size=8.5, color=MUTED)

        if text.startswith("H1.") or text.startswith("H2."):
            p.paragraph_format.left_indent = Cm(0.7)
            p.paragraph_format.right_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(7)
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), PALE)
            p_pr.append(shd)

        if style_name == "Caption":
            p.paragraph_format.keep_with_next = True

        # Keep standalone table titles with the following table.
        if text.startswith("Tabela ") and len(text) < 120:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                style_run(run, size=9.5, color=NAVY, bold=True)


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        table.autofit = True
        table.alignment = 0
        table.style = get_style(doc, "Table")
        if not table.rows:
            continue
        set_repeat_table_header(table.rows[0])
        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Cm(0.55)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                set_cell_shading(cell, NAVY if row_index == 0 else (PALE if row_index % 2 == 0 else WHITE))
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.line_spacing = 1.0
                    for run in p.runs:
                        style_run(
                            run,
                            size=8.2,
                            color=WHITE if row_index == 0 else TEXT,
                            bold=True if row_index == 0 else None,
                        )


def resize_images(doc: Document) -> None:
    max_width = Cm(15.8)
    for shape in doc.inline_shapes:
        if shape.width > max_width:
            ratio = max_width / shape.width
            shape.width = max_width
            shape.height = int(shape.height * ratio)

    # Image paragraphs stay with their captions and avoid orphaned graphics.
    for p in doc.paragraphs:
        if p._p.xpath(".//w:drawing"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(2)


def add_bottom_border_to_title(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.style and p.style.name == "Title":
            p_pr = p._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "14")
            bottom.set(qn("w:space"), "8")
            bottom.set(qn("w:color"), TEAL)
            borders.append(bottom)
            p_pr.append(borders)
            break


def set_document_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Coordenação sem recursos? Efeitos da adesão municipal ao SISAN"
    props.subject = "Avaliação por matching das compras da agricultura familiar no PNAE"
    props.author = "Osmar Junior"
    props.keywords = "SISAN; PNAE; matching; avaliação de políticas públicas"
    props.comments = "Versão gerada pelo suplemento reproduzível em 2026-08-28."


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    resource_path = os.pathsep.join([str(ROOT / "manuscript"), str(ROOT)])
    subprocess.run(
        [
            "pandoc",
            str(SOURCE),
            "--from=markdown+tex_math_dollars",
            "--to=docx",
            "--standalone",
            f"--resource-path={resource_path}",
            "--output",
            str(OUTPUT),
        ],
        cwd=ROOT,
        check=True,
    )

    doc = Document(OUTPUT)
    format_styles(doc)
    format_sections(doc)
    format_title_block(doc)
    add_bottom_border_to_title(doc)
    format_paragraphs(doc)
    format_tables(doc)
    resize_images(doc)
    set_document_metadata(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
